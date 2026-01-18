"""Pytest fixtures for PIKA tests."""

import os
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient

# Set test environment variables before importing app
os.environ["PIKA_SESSION_SECRET"] = "test-secret-for-testing-only"
os.environ["DOCUMENTS_DIR"] = tempfile.mkdtemp()
os.environ["CHROMA_PERSIST_DIR"] = tempfile.mkdtemp()
os.environ["AUDIT_LOG_PATH"] = tempfile.mktemp(suffix=".log")


@pytest.fixture(scope="session")
def temp_dirs():
    """Create temporary directories for tests."""
    docs_dir = tempfile.mkdtemp()
    data_dir = tempfile.mkdtemp()
    yield {"docs": docs_dir, "data": data_dir}
    # Cleanup handled by OS for temp dirs


@pytest.fixture
def test_client():
    """Create a test client for the FastAPI app."""
    # Import here to ensure env vars are set
    from pika.main import app

    with TestClient(app) as client:
        yield client


@pytest.fixture
def mock_ollama_client():
    """Create a mock Ollama client."""
    mock = MagicMock()
    mock.health_check = AsyncMock(return_value=True)
    mock.list_models = AsyncMock(return_value=[])
    mock.generate = AsyncMock(return_value="Test response")
    mock.model = "test-model"
    return mock


@pytest.fixture
def mock_rag_engine():
    """Create a mock RAG engine."""
    from pika.services.rag import IndexStats, QueryResult, Confidence, Source

    mock = MagicMock()
    mock.get_stats = MagicMock(return_value=IndexStats(
        total_documents=5,
        total_chunks=50,
        collection_name="test_collection",
    ))
    mock.index_documents = MagicMock(return_value=IndexStats(
        total_documents=5,
        total_chunks=50,
        collection_name="test_collection",
    ))
    mock.query = AsyncMock(return_value=QueryResult(
        answer="Test answer",
        sources=[
            Source(
                filename="test.txt",
                chunk_index=0,
                content="Test content",
                similarity=0.9,
            )
        ],
        confidence=Confidence.HIGH,
    ))
    return mock


@pytest.fixture
def sample_document(temp_dirs):
    """Create a sample document for testing."""
    docs_dir = Path(temp_dirs["docs"])
    docs_dir.mkdir(parents=True, exist_ok=True)

    doc_path = docs_dir / "test_document.txt"
    doc_path.write_text("This is a test document with some content for testing purposes.")

    yield doc_path

    # Cleanup
    if doc_path.exists():
        doc_path.unlink()


@pytest.fixture
def auth_headers():
    """Generate headers with a valid API key for testing."""
    return {"X-API-Key": "test-api-key"}


@pytest.fixture(autouse=True)
def reset_singletons():
    """Reset singleton instances between tests."""
    # Reset app config singleton
    import pika.services.app_config as app_config_module
    app_config_module._app_config = None

    # Reset audit logger singleton
    import pika.services.audit as audit_module
    audit_module._audit_logger = None

    # Reset RAG engine singleton
    import pika.services.rag as rag_module
    rag_module._rag_engine = None

    # Reset sessions
    import pika.api.web as web_module
    with web_module._sessions_lock:
        web_module._sessions.clear()
    with web_module._csrf_lock:
        web_module._csrf_tokens.clear()

    yield


@pytest.fixture
def sample_pdf_file(temp_dirs):
    """Create a sample PDF file for testing."""
    from pypdf import PdfWriter
    from io import BytesIO

    docs_dir = Path(temp_dirs["docs"])
    docs_dir.mkdir(parents=True, exist_ok=True)

    # Create a minimal PDF with text
    pdf_path = docs_dir / "test_document.pdf"

    writer = PdfWriter()
    # Add a blank page and annotate it with text
    page = writer.add_blank_page(width=612, height=792)

    # Create PDF with embedded text using a simpler approach
    # Write a minimal PDF manually that contains extractable text
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 84 >>
stream
BT
/F1 12 Tf
100 700 Td
(This is a test PDF document for PIKA testing.) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000400 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
477
%%EOF"""

    pdf_path.write_bytes(pdf_content)

    yield pdf_path

    if pdf_path.exists():
        pdf_path.unlink()


@pytest.fixture
def sample_docx_file(temp_dirs):
    """Create a sample DOCX file for testing."""
    from docx import Document

    docs_dir = Path(temp_dirs["docs"])
    docs_dir.mkdir(parents=True, exist_ok=True)

    docx_path = docs_dir / "test_document.docx"

    # Create a Word document
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test DOCX document for PIKA testing.")
    doc.add_paragraph("It contains multiple paragraphs to test extraction.")
    doc.save(str(docx_path))

    yield docx_path

    if docx_path.exists():
        docx_path.unlink()
